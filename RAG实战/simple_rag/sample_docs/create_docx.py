import os

sample_docs_dir = os.path.dirname(__file__)

try:
    from docx import Document
    doc = Document()
    doc.add_heading("RAG 系统技术文档", 0)
    doc.add_heading("1. 概述", level=1)
    doc.add_paragraph("RAG（检索增强生成）是一种结合检索系统和生成模型的技术，能够从外部知识库检索相关信息来增强大语言模型的能力。")
    doc.add_heading("2. 核心组件", level=1)
    doc.add_paragraph("• 文档解析器：支持PDF、Word、Excel等格式")
    doc.add_paragraph("• 文本分块器：将长文档切分成小片段")
    doc.add_paragraph("• 向量索引：使用FAISS构建高效索引")
    doc.add_paragraph("• 混合检索：结合BM25和向量检索")
    doc.add_heading("3. 应用场景", level=1)
    doc.add_paragraph("适用于企业知识库、智能客服、技术文档问答等场景。")
    doc.save(os.path.join(sample_docs_dir, "技术文档.docx"))
    print("✓ Word文档创建成功")
except Exception as e:
    print(f"创建Word文档失败: {e}")
