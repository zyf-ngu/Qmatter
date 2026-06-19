#python-docx解析
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.shape import CT_Picture
from docx.oxml.table import CT_Tbl
import pytesseract
from PIL import Image

docx_path = 'example.docx'


def process_docx_with_runs(docx_path):
    doc = Document(docx_path)
    output = []
# 遍历所有段落
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:
            output.append(('text', para_text))
        for run in para.runs:
            if run.text:
                output.append(run.text)
        print('run-output', output)
        # for run in para.runs:
        # # 检测 Run 中的图片
        #     for elem in run._element.xpath('.//pic:pic',
        #                                namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
        #         image_part = None
        # # 获取图片资源 ID
        #     blip = elem.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})[0]
        #     image_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        # # 通过文档的 Relationships 获取图片路径
        #     image_part = doc.part.related_parts[image_id]
        # # 保存图片并 OCR
        #     with open('temp_image.png', 'wb') as f:
        #         f.write(image_part.blob)
        #     ocr_text = pytesseract.image_to_string(Image.open('temp_image.png'))
        #     output.append(('image_ocr', ocr_text))
    print('para-output',output)
    for idx, shape in enumerate(doc.inline_shapes):
        if shape.type == 1:  # 1 代表图片（WD_INLINE_SHAPE_TYPE.PICTURE）
        # 获取图片原始数据（blob）
            image_blob =shape.image.blob
        # 保存到本地
            with open(f"extracted_image_{idx}.png", "wb") as f:
                f.write(image_blob)
    # 遍历段落内的 Run


      # 处理表格
    for table in doc.tables:
         table_data = []
         for row in table.rows:
              row_data = [cell.text.strip() for cell in row.cells]
              table_data.append(''.join(row_data))
         output.append(('table', '\n'.join(table_data)))
    print('table-output', output)
    return output


result=process_docx_with_runs(docx_path=docx_path)



#lxml解析
# 解压.docx并读取核心XML
# 首先解压.docx，读取两个关键文件：
# word / document.xml：文档主体内容（段落、表格、图片引用）；
# word / _rels / document.xml.rels：图片等资源的关系映射（将XML中的r: embedID映射到实际图片路径）。
import zipfile
from lxml import etree
import os


def load_docx_content(docx_path, output_dir='extracted_images'):
    """解压 docx 并提取段落、表格和图片，返回 (paragraphs, tables, image_paths)"""
    with zipfile.ZipFile(docx_path, 'r') as z:
        # 1. 读取 XML
        document_xml = z.read('word/document.xml')
        rels_xml = z.read('word/_rels/document.xml.rels')

        # 2. 定义命名空间（document.xml 用）
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        # 3. 解析 .rels 文件（使用正确的命名空间）
        rels_root = etree.fromstring(rels_xml)
        rels_ns = {'rel': 'http://schemas.openxmlformats.org/package/2006/relationships'}
        rels_map = {}
        for rel in rels_root.xpath('//rel:Relationship', namespaces=rels_ns):
            rel_id = rel.get('Id')
            target = rel.get('Target')
            rels_map[rel_id] = target

        # 4. 解析 document.xml
        doc_root = etree.fromstring(document_xml)

        # 5. 提取段落
        paragraphs = []
        for p in doc_root.findall('.//w:p', namespaces):
            text_parts = [t.text for t in p.findall('.//w:t', namespaces) if t.text]
            paragraphs.append(''.join(text_parts))

        # 6. 提取表格
        tables = []
        for tbl in doc_root.findall('.//w:tbl', namespaces):
            table_data = []
            for tr in tbl.findall('.//w:tr', namespaces):
                row = []
                for tc in tr.findall('.//w:tc', namespaces):
                    cell_text = ''.join(
                        t.text for t in tc.findall('.//w:t', namespaces) if t.text
                    )
                    row.append(cell_text)
                table_data.append(row)
            tables.append(table_data)

        # 7. 提取图片
        os.makedirs(output_dir, exist_ok=True)
        image_paths = []
        for blip in doc_root.findall('.//a:blip', namespaces):
            rel_id = blip.get(f'{{{namespaces["r"]}}}embed')
            if not rel_id:
                continue
            target_path = rels_map.get(rel_id)
            if not target_path:
                print(f"警告：找不到关系 {rel_id}")
                continue
            # 补全路径（.rels 中的 target 通常是 'media/image1.png'）
            full_path = f'word/{target_path}'
            try:
                img_data = z.read(full_path)
                ext = target_path.split('.')[-1]
                img_filename = f'image_{len(image_paths)}.{ext}'
                img_save_path = os.path.join(output_dir, img_filename)
                with open(img_save_path, 'wb') as f:
                    f.write(img_data)
                image_paths.append(img_save_path)
                print(f"已保存图片: {img_save_path}")
            except KeyError:
                print(f"未找到图片文件: {full_path}")

        return paragraphs, tables, image_paths


paragraphs, tables, images = load_docx_content(docx_path)

print("=== 段落 ===")
for i, para in enumerate(paragraphs, 1):
    print(f"{i}: {para}")

print("\n=== 表格 ===")
for i, table in enumerate(tables, 1):
    print(f"表格 {i}:", table)

print(f"\n=== 图片 ===")
for img in images:
    print(img)

