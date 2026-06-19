from typing import List
import os


class DocxParser:
    """Word 文档解析器"""
    
    def parse(self, file_path: str) -> List[str]:
        try:
            from docx import Document
            doc = Document(file_path)
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        texts.append(row_text)
            return texts
        except ImportError:
            try:
                import zipfile
                from lxml import etree
                with zipfile.ZipFile(file_path, 'r') as z:
                    document_xml = z.read('word/document.xml')
                    namespaces = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    }
                    doc_root = etree.fromstring(document_xml)
                    paragraphs = []
                    for p in doc_root.findall('.//w:p', namespaces):
                        text_parts = [t.text for t in p.findall('.//w:t', namespaces) if t.text]
                        if text_parts:
                            paragraphs.append(''.join(text_parts))
                    return paragraphs
            except Exception as e:
                print(f"DOCX 解析失败: {e}")
                return []


if __name__ == "__main__":
    parser = DocxParser()
    # result = parser.parse("example.docx")
    # print(result)
